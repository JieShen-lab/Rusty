from __future__ import annotations

from pathlib import Path

from docx import Document

from rusty.models import ParsedBook, ParsedChapter


def parse_docx(path: str | Path) -> ParsedBook:
    source_path = Path(path)
    document = Document(str(source_path))
    properties = document.core_properties
    title = (properties.title or "").strip() or source_path.stem
    author = (properties.author or "").strip() or None
    description = (properties.comments or "").strip() or None
    chapters = _split_docx_chapters(document, fallback_title=title)

    return ParsedBook(
        title=title,
        author=author,
        language=None,
        source_path=source_path,
        source_format="docx",
        source_encoding=None,
        chapters=chapters,
        description=description,
        metadata={
            "category": properties.category,
            "keywords": properties.keywords,
            "subject": properties.subject,
        },
    )


def _split_docx_chapters(document: Document, fallback_title: str) -> list[ParsedChapter]:
    chapters: list[ParsedChapter] = []
    current_title: str | None = None
    current_lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if _is_heading(paragraph):
            if current_title is not None:
                chapters.append(
                    ParsedChapter(
                        index=len(chapters) + 1,
                        title=current_title,
                        text="\n".join(current_lines).strip(),
                    )
                )
            current_title = text
            current_lines = []
            continue

        current_lines.append(text)

    if current_title is not None:
        chapters.append(
            ParsedChapter(
                index=len(chapters) + 1,
                title=current_title,
                text="\n".join(current_lines).strip(),
            )
        )

    if chapters:
        return chapters

    body = "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
    return [ParsedChapter(index=1, title=fallback_title, text=body)]


def _is_heading(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    normalized = style_name.lower()
    return normalized.startswith("heading") or normalized.startswith("标题")

