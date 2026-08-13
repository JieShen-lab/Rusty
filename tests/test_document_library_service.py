from __future__ import annotations

import sqlite3
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from tests.support import initialized_database

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document
from ebooklib import epub

from rusty.db import session
from rusty.models import ParsedBook, ParsedChapter, count_text_units
from rusty.services.project_service import ProjectService
from rusty.services.document_library_service import DocumentLibraryService, DraftConflictError


class DocumentLibraryServiceTests(unittest.TestCase):
    def test_chapter_ordinals_are_derived_and_exports_compose_them_once(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "titles.txt"
            source.write_text("第一章 阿尔法为父亲\n\n正文一。\n", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            first = service.list_chapters(document.id)[0]

            created = service.create_chapter(
                document.id,
                title="",
                text="正文二。",
                position="after",
                anchor_chapter_id=first.id,
            )
            chapters = service.list_chapters(document.id)
            exported = service.export_document(document.id, "txt", root / "export.txt").read_text(encoding="utf-8")

            self.assertIsNotNone(created.created_chapter_id)
            self.assertEqual([1, 2], [chapter.index for chapter in chapters])
            self.assertEqual(["阿尔法为父亲", ""], [chapter.title for chapter in chapters])
            self.assertEqual(1, exported.count("第一章 阿尔法为父亲"))
            self.assertEqual(1, exported.count("第二章"))
            self.assertNotIn("第二章 第二章", exported)

    def test_draft_autosave_does_not_create_revision_and_manual_commit_creates_one(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "draft.txt"
            source.write_text("第一章\n\n原正文。\n\n第二章\n\n后续正文。\n", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            chapter = service.list_chapters(document.id)[0]
            content = service.get_content(document.id, chapter.id)
            revision_count = len(service.list_revisions(document.id))

            draft = service.save_draft(
                document.id,
                chapter_id=chapter.id,
                base_revision_id=content.revision_id,
                title="第一章（修订）",
                text="草稿正文🙂。",
            )

            self.assertEqual(revision_count, len(service.list_revisions(document.id))
            )
            self.assertEqual("草稿正文🙂。", service.get_draft(document.id, chapter.id).text)
            committed = service.commit_draft(document.id, chapter.id)
            self.assertEqual(revision_count + 1, len(service.list_revisions(document.id)))
            self.assertEqual("manual_edit", committed.revision.revision_type)
            self.assertIsNone(service.get_draft(document.id, chapter.id))
            current_chapters = service.list_chapters(document.id)
            renamed = next(item for item in current_chapters if item.index == 1)
            self.assertEqual("（修订）", renamed.title)
            self.assertEqual("草稿正文🙂。\n\n", service.get_content(document.id, renamed.id).body_text)
            self.assertEqual(draft.base_revision_id, committed.revision.parent_revision_id)

    def test_stale_chapter_draft_conflicts_after_revision_changes(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "conflict.txt"
            source.write_text("第一章\n\n正文。\n", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            chapter = service.list_chapters(document.id)[0]
            content = service.get_content(document.id, chapter.id)
            service.save_draft(
                document.id,
                chapter_id=chapter.id,
                base_revision_id=content.revision_id,
                title=content.title,
                text="尚未提交",
            )
            service.save_content(document.id, text="全文替换")

            with self.assertRaises(DraftConflictError):
                service.commit_draft(document.id, chapter.id)
            self.assertEqual("尚未提交", service.get_draft(document.id, chapter.id).text)

    def test_chapter_title_body_offsets_and_unicode_counts_stay_consistent(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source_text = "第一章\n\n中 A，。\u3000\n🙂\n\n第二章\n\n尾声。\n"
            source = root / "counts.txt"
            source.write_text(source_text, encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            first, second = service.list_chapters(document.id)
            first_content = service.get_content(document.id, first.id)
            old_second_start = second.start_offset

            self.assertEqual("", first_content.title)
            self.assertNotIn("第一章", first_content.body_text)
            self.assertGreater(first_content.body_start_offset, first_content.section_start_offset)
            self.assertEqual(count_text_units(source_text), document.word_count)

            service.save_draft(
                document.id,
                chapter_id=first.id,
                base_revision_id=first_content.revision_id,
                title="新标题",
                text="加长正文 ABC🙂。",
            )
            service.commit_draft(document.id, first.id)
            updated = service.list_chapters(document.id)
            updated_first, updated_second = updated
            updated_content = service.get_content(document.id, updated_first.id)
            stored = Path(service.list_revisions(document.id)[0].storage_path).read_text(encoding="utf-8")

            self.assertEqual("新标题", updated_first.title)
            self.assertTrue(stored[updated_first.start_offset:updated_first.end_offset].lstrip("\n").startswith("第一章 新标题\n\n"))
            self.assertGreater(updated_second.start_offset, old_second_start)
            self.assertEqual(count_text_units(stored), service.list_documents()[0].word_count)

    def test_regex_split_creates_revision_and_old_revision_can_be_restored(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "split.txt"
            source.write_text("第一章\n正文\n第二章\n正文\n", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            original = service.list_revisions(document.id)[0]
            preview = service.preview_regex_split(document.id, r"^第.+章$")

            saved = service.apply_regex_split(document.id, r"^第.+章$", preview.preview_token)

            self.assertEqual(2, len(saved))
            current = service.list_revisions(document.id)[0]
            self.assertEqual("split_regex", current.revision_type)
            self.assertEqual(2, len(service.list_revisions(document.id)))
            restored = service.activate_revision(document.id, original.id)
            self.assertEqual(original.storage_path, restored.storage_path)

    def test_explicit_split_writes_final_boundaries_without_reparsing_text(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "explicit-split.txt"
            text = "第一章\n正文一\n第二章\n正文二\n"
            source.write_text(text, encoding="utf-8")
            service = DocumentLibraryService(
                initialized_database(root / "rusty.db"), root / "library"
            )
            document = service.import_document(source).document
            original = service.list_revisions(document.id)[0]
            second_start = text.index("第二章")
            boundaries = [
                {"title": "一", "start_offset": 0, "end_offset": second_start},
                {
                    "title": "二",
                    "start_offset": second_start,
                    "end_offset": len(text),
                },
            ]

            with mock.patch(
                "rusty.services.document_library_service.parse_txt",
                side_effect=AssertionError("explicit boundaries must not parse"),
            ):
                revision, chapters = service.apply_split_boundaries(
                    document.id,
                    source_revision_id=original.id,
                    boundaries=boundaries,
                    revision_type="split_manual",
                )

            self.assertEqual(["一", "二"], [chapter.title for chapter in chapters])
            self.assertEqual(
                [(0, second_start), (second_start, len(text))],
                [(chapter.start_offset, chapter.end_offset) for chapter in chapters],
            )
            with session(service.database_path) as connection:
                stored = connection.execute(
                    "SELECT COUNT(*) FROM library_document_chapters WHERE revision_id = ?",
                    (revision.id,),
                ).fetchone()[0]
            self.assertEqual(2, stored)

    def test_split_failure_rolls_back_revision_head_and_projection(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "split-rollback.txt"
            text = "第一章\n正文一\n第二章\n正文二\n"
            source.write_text(text, encoding="utf-8")
            service = DocumentLibraryService(
                initialized_database(root / "rusty.db"), root / "library"
            )
            document = service.import_document(source).document
            original = service.list_revisions(document.id)[0]
            original_chapters = service.list_chapters(document.id)
            second_start = text.index("第二章")

            with (
                mock.patch(
                    "rusty.services.document_library_service.count_text_units",
                    side_effect=RuntimeError("injected split failure"),
                ),
                self.assertRaisesRegex(RuntimeError, "injected split failure"),
            ):
                service.apply_split_boundaries(
                    document.id,
                    source_revision_id=original.id,
                    boundaries=[
                        {"title": "一", "start_offset": 0, "end_offset": second_start},
                        {
                            "title": "二",
                            "start_offset": second_start,
                            "end_offset": len(text),
                        },
                    ],
                    revision_type="split_manual",
                )

            self.assertEqual(original.id, service._ensure_initial_revision(document.id).id)
            self.assertEqual(1, len(service.list_revisions(document.id)))
            self.assertEqual(
                [(item.title, item.start_offset, item.end_offset) for item in original_chapters],
                [
                    (item.title, item.start_offset, item.end_offset)
                    for item in service.list_chapters(document.id)
                ],
            )

    def test_draft_commit_failure_rolls_back_revision_and_preserves_exact_draft(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "draft-rollback.txt"
            source.write_text("第一章\n正文\n", encoding="utf-8")
            service = DocumentLibraryService(
                initialized_database(root / "rusty.db"), root / "library"
            )
            document = service.import_document(source).document
            content = service.get_content(document.id)
            draft = service.save_draft(
                document.id,
                base_revision_id=content.revision_id,
                title=content.title,
                text="替换正文",
            )

            with (
                mock.patch.object(
                    service,
                    "_insert_chapters",
                    side_effect=RuntimeError("injected draft failure"),
                ),
                self.assertRaisesRegex(RuntimeError, "injected draft failure"),
            ):
                service.commit_draft(document.id)

            self.assertEqual(1, len(service.list_revisions(document.id)))
            self.assertEqual(draft.id, service.get_draft(document.id).id)
            self.assertEqual(
                content.revision_id, service._ensure_initial_revision(document.id).id
            )

    def test_draft_title_rolls_back_when_exact_draft_consumption_fails(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "draft-title-rollback.txt"
            source.write_text("第一章\n正文\n", encoding="utf-8")
            service = DocumentLibraryService(
                initialized_database(root / "rusty.db"), root / "library"
            )
            document = service.import_document(source).document
            content = service.get_content(document.id)
            draft = service.save_draft(
                document.id,
                base_revision_id=content.revision_id,
                title="新书名",
                text="替换正文",
            )
            with session(service.database_path) as connection:
                connection.execute(
                    """
                    CREATE TRIGGER fail_exact_draft_delete
                    BEFORE DELETE ON library_document_drafts
                    BEGIN
                        SELECT RAISE(ABORT, 'injected draft delete failure');
                    END
                    """
                )

            with self.assertRaisesRegex(sqlite3.IntegrityError, "injected draft delete failure"):
                service.commit_draft(document.id)

            self.assertEqual(1, len(service.list_revisions(document.id)))
            self.assertEqual(draft.id, service.get_draft(document.id).id)
            self.assertEqual(document.title, service._get_document(document.id).title)
            self.assertEqual(
                content.revision_id, service._ensure_initial_revision(document.id).id
            )

    def test_committing_one_chapter_draft_does_not_delete_another(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "two-drafts.txt"
            source.write_text("第一章\n正文一\n第二章\n正文二\n", encoding="utf-8")
            service = DocumentLibraryService(
                initialized_database(root / "rusty.db"), root / "library"
            )
            document = service.import_document(source).document
            first, second = service.list_chapters(document.id)
            first_content = service.get_content(document.id, first.id)
            second_content = service.get_content(document.id, second.id)
            service.save_draft(
                document.id,
                chapter_id=first.id,
                base_revision_id=first_content.revision_id,
                title=first_content.title,
                text="第一章新正文",
            )
            other = service.save_draft(
                document.id,
                chapter_id=second.id,
                base_revision_id=second_content.revision_id,
                title=second_content.title,
                text="第二章未提交正文",
            )

            service.commit_draft(document.id, first.id)

            self.assertIsNone(service.get_draft(document.id, first.id))
            self.assertEqual(other.id, service.get_draft(document.id, second.id).id)

    def test_categories_are_many_to_many_and_independent_from_tags(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "classified.txt"
            source.write_text("正文", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            first = service.create_category("参考")
            second = service.create_category("待整理")
            tag = service.create_tag("长篇")
            second = service.rename_category(second.id, "归档")
            with self.assertRaisesRegex(ValueError, "已存在"):
                service.rename_category(second.id, "参考")

            service.set_document_category(document.id, first.id, True)
            service.set_document_category(document.id, second.id, True)
            service.set_document_tag(document.id, tag.id, True)
            assigned = service.list_documents()[0]

            self.assertEqual({first.id, second.id}, set(assigned.category_ids))
            self.assertEqual({"参考", "归档"}, set(assigned.categories))
            self.assertEqual(["长篇"], assigned.tags)
            service.delete_category(first.id)
            remaining = service.list_documents()[0]
            self.assertEqual([second.id], remaining.category_ids)
            self.assertEqual(["归档"], remaining.categories)
            self.assertEqual(["长篇"], remaining.tags)
            self.assertTrue(Path(remaining.storage_path).is_file())

    def test_project_document_uses_relation_without_creating_legacy_tag(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "project-source.txt"
            source.write_text("工程正文", encoding="utf-8")
            database = initialized_database(root / "rusty.db")
            service = DocumentLibraryService(database, root / "library")
            ordinary = service.import_document(source).document
            self.assertFalse(ordinary.is_project_document)
            project_id = ProjectService(database).create_project(
                ParsedBook(
                    title="工程",
                    author="",
                    language="zh",
                    source_path=source,
                    source_format="txt",
                    source_encoding="utf-8",
                    chapters=[ParsedChapter(index=1, title="第一章", text="工程正文")],
                ),
                root / "workspace",
            )

            linked = service.ensure_project_document(project_id, source)

            self.assertTrue(linked.is_project_document)
            self.assertEqual([project_id], linked.project_ids)
            self.assertNotIn("工程", linked.tags)
            self.assertNotIn("工程", [tag.name for tag in service.list_tags()])
            with self.assertRaisesRegex(ValueError, "仅支持阅读"):
                service.save_content(linked.id, text="不应写回工程")
            project_chapter = ProjectService(database).list_chapters(project_id)[0]
            ProjectService(database).save_chapter_rewrite(project_chapter.id, "工程最新保存正文")
            synced = service.sync_project_document(project_id)
            self.assertIsNotNone(synced)
            synced_chapter = service.list_chapters(linked.id)[0]
            self.assertIn("工程最新保存正文", service.get_content(linked.id, synced_chapter.id).body_text)
            self.assertEqual("project_sync", service.list_revisions(linked.id)[0].revision_type)
            ordinary_source = root / "ordinary.txt"
            ordinary_source.write_text("普通导入正文", encoding="utf-8")
            ordinary_document = service.import_document(ordinary_source).document
            all_documents = service.list_documents()
            self.assertEqual({linked.id, ordinary_document.id}, {item.id for item in all_documents})
            self.assertEqual([linked.id], [item.id for item in all_documents if item.is_project_document])
            self.assertEqual([ordinary_document.id], [item.id for item in service.list_recent_imports()])
            with session(database) as connection:
                self.assertEqual(
                    linked.id,
                    connection.execute(
                        "SELECT document_id FROM project_documents WHERE project_id = ?",
                        (project_id,),
                    ).fetchone()[0],
                )

    def test_cursor_split_preserves_order_offsets_and_creates_revision(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "cursor-split.txt"
            source.write_text("第一章 原标题\n\nAAAAAABBBBBB\n\n第二章 后续\n\nCCCCCC\n", encoding="utf-8")
            database = initialized_database(root / "rusty.db")
            service = DocumentLibraryService(database, root / "library")
            document = service.import_document(source).document
            first, second = service.list_chapters(document.id)
            before_revisions = len(service.list_revisions(document.id))

            result = service.split_chapter_at_cursor(
                document.id,
                chapter_id=first.id,
                cursor_offset=6,
                next_title="新章节",
            )

            chapters = service.list_chapters(document.id)
            self.assertEqual(before_revisions + 1, len(service.list_revisions(document.id)))
            self.assertEqual("split_cursor", result.revision.revision_type)
            self.assertEqual(["原标题", "新章节", "后续"], [chapter.title for chapter in chapters])
            self.assertEqual([1, 2, 3], [chapter.index for chapter in chapters])
            self.assertEqual("AAAAAA", service.get_content(document.id, chapters[0].id).body_text)
            self.assertTrue(service.get_content(document.id, chapters[1].id).body_text.startswith("BBBBBB"))
            self.assertEqual(second.title, chapters[2].title)
            self.assertEqual(
                chapters[0].end_offset,
                chapters[1].start_offset,
            )
            self.assertEqual(
                chapters[1].end_offset,
                chapters[2].start_offset,
            )

    def test_prompt_cleanup_creates_cleanup_revision_and_preserves_chapter_order(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "prompt-cleanup.txt"
            source.write_text("第一章\n\n正文一。\n\n第二章\n\n正文二。\n", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            first = service.list_chapters(document.id)[0]

            result = service.apply_prompt_cleanup(
                document.id,
                chapter_id=first.id,
                title=first.title,
                cleaned_text="正文一。\n",
                prompt="只整理空白",
            )

            self.assertEqual("cleanup_ai", result.revision.revision_type)
            self.assertEqual(2, len(service.list_chapters(document.id)))
            self.assertEqual(["", ""], [chapter.title for chapter in service.list_chapters(document.id)])

    def test_default_cleanup_template_versions_text_and_can_restore_import(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "排版.txt"
            source.write_text("  第一章 风起\n第一段。\n\n　第二段。\n", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            imported = service.import_document(source)
            template = service.list_processing_templates()[0]

            cleaned = service.apply_cleanup(imported.document.id, template.id)
            repeated = service.apply_cleanup(imported.document.id, template.id)
            revisions = service.list_revisions(imported.document.id)
            cleaned_text = Path(cleaned.document.storage_path).read_text(encoding="utf-8")

            self.assertTrue(template.is_default)
            self.assertEqual(0, template.settings["chapter_indent"])
            self.assertEqual(2, template.settings["paragraph_indent"])
            self.assertEqual("第一章 风起\n\n　　第一段。\n\n　　第二段。\n", cleaned_text)
            self.assertTrue(cleaned.created)
            self.assertFalse(repeated.created)
            self.assertEqual(2, len(revisions))
            self.assertEqual("processed", cleaned.document.status)

            restored = service.activate_revision(imported.document.id, revisions[-1].id)
            self.assertEqual("imported", restored.status)
            self.assertEqual(imported.document.storage_path, restored.storage_path)

    def test_txt_is_normalized_to_utf8_and_duplicate_content_is_reused(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "原稿.txt"
            source.write_bytes("第一章\r\n正文。".encode("gb18030"))
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")

            first = service.import_document(source)
            second = service.import_document(source)
            renamed = service.update_document_metadata(
                first.document.id,
                title="新的书名",
                author="作者甲",
            )

            stored_path = Path(first.document.storage_path)
            self.assertTrue(first.created)
            self.assertFalse(second.created)
            self.assertEqual(first.document.id, second.document.id)
            self.assertEqual(".txt", stored_path.suffix)
            self.assertEqual("第一章\n正文。\n", stored_path.read_text(encoding="utf-8"))
            self.assertEqual("txt", first.document.source_format)
            self.assertEqual("新的书名", renamed.title)
            self.assertEqual("作者甲", renamed.author)
            self.assertEqual(1, len(service.list_documents()))

    def test_tags_chapters_migration_and_exports_are_functional(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "长篇.txt"
            source.write_text(
                "第一章 风起\n第一段。\n\n第二章 归途\n第二段。\n",
                encoding="utf-8",
            )
            database_path = initialized_database(root / "rusty.db")
            original_library = root / "library"
            migrated_library = root / "migrated-library"
            service = DocumentLibraryService(database_path, original_library)

            imported = service.import_document(source)
            tag = service.create_tag("参考资料")
            assigned = service.set_document_tag(imported.document.id, tag.id, True)
            chapters = service.list_chapters(imported.document.id)
            full_content = service.get_content(imported.document.id)
            chapter_content = service.get_content(imported.document.id, chapters[1].id)
            reordered = service.reorder_chapters(
                imported.document.id,
                [chapters[1].id, chapters[0].id],
            )
            original_paths = [Path(item.storage_path) for item in service.list_revisions(imported.document.id)]

            migrated = service.migrate_library_path(migrated_library)
            txt_output = service.export_document(imported.document.id, "txt", root / "export.txt")
            epub_output = service.export_document(imported.document.id, "epub", root / "export.epub")
            restarted = DocumentLibraryService(database_path)

            self.assertEqual(["参考资料"], assigned.tags)
            self.assertEqual(1, service.list_tags()[0].resource_count)
            self.assertEqual(["风起", "归途"], [item.title for item in chapters])
            self.assertIn("第一章 风起", full_content.text)
            self.assertEqual("归途", chapter_content.title)
            self.assertIn("第二段。", chapter_content.text)
            self.assertEqual(["归途", "风起"], [item.title for item in reordered])
            self.assertEqual(migrated_library.resolve(), migrated)
            self.assertEqual(migrated_library.resolve(), restarted.get_library_path().resolve())
            self.assertTrue(all(not path.exists() for path in original_paths))
            self.assertTrue(all(Path(item.storage_path).parent == migrated_library.resolve() for item in service.list_revisions(imported.document.id)))
            exported_text = txt_output.read_text(encoding="utf-8")
            self.assertIn("第一章 归途", exported_text)
            self.assertIn("第二章 风起", exported_text)
            self.assertLess(
                exported_text.index("第一章 归途"),
                exported_text.index("第二章 风起"),
            )
            self.assertTrue(epub_output.is_file())
            self.assertGreater(epub_output.stat().st_size, 0)
            service.delete_document(imported.document.id)
            self.assertEqual([], service.list_documents())

    def test_docx_is_flattened_to_txt_with_metadata(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "word.docx"
            document = Document()
            document.core_properties.title = "Word 示例"
            document.core_properties.author = "作者甲"
            document.add_heading("第一章", level=1)
            document.add_paragraph("第一段正文。")
            document.add_heading("第二章", level=1)
            document.add_paragraph("第二段正文。")
            document.save(source)
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")

            result = service.import_document(source)
            stored_text = Path(result.document.storage_path).read_text(encoding="utf-8")

            self.assertEqual("Word 示例", result.document.title)
            self.assertEqual("作者甲", result.document.author)
            self.assertEqual("docx", result.document.source_format)
            self.assertEqual(2, result.document.chapter_count)
            self.assertIn("第一章\n\n第一段正文。", stored_text)
            self.assertIn("第二章\n\n第二段正文。", stored_text)
            self.assertEqual(".txt", Path(result.document.storage_path).suffix)

    def test_epub_is_flattened_to_txt_with_metadata(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "book.epub"
            _write_sample_epub(source)
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")

            result = service.import_document(source)
            stored_text = Path(result.document.storage_path).read_text(encoding="utf-8")

            self.assertEqual("EPUB 示例", result.document.title)
            self.assertEqual("作者乙", result.document.author)
            self.assertEqual("epub", result.document.source_format)
            self.assertIn("开篇\n\n第一段。", stored_text)
            self.assertEqual(".txt", Path(result.document.storage_path).suffix)

    def test_volume_hierarchy_survives_chapter_edit_and_revision_restore(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            source = root / "volumes.txt"
            source.write_text(
                "第七卷 雨夜\n\n第787章 雨夜\n正文一。\n\n第788章 风声\n正文二。\n",
                encoding="utf-8",
            )
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            document = service.import_document(source).document
            original_revision = service.list_revisions(document.id)[0]
            directory_before = service.get_directory(document.id)

            first_chapter = directory_before.volumes[0][1][0]
            edit = service.rename_chapter(document.id, first_chapter.id, "第787章 新雨")
            directory_after = service.get_directory(document.id)

            self.assertEqual(1, len(directory_after.volumes))
            self.assertEqual("第七卷 雨夜", directory_after.volumes[0][0].title)
            self.assertEqual(
                ["新雨", "风声"],
                [chapter.title for chapter in directory_after.volumes[0][1]],
            )
            self.assertTrue(all(chapter.volume_id == directory_after.volumes[0][0].id for chapter in directory_after.volumes[0][1]))

            renamed_volume = service.rename_volume(
                document.id,
                directory_after.volumes[0][0].id,
                "雨夜篇",
            )
            service.apply_cleanup(
                document.id,
                service.list_processing_templates()[0].id,
            )
            after_cleanup = service.get_directory(document.id)
            self.assertEqual("雨夜篇", after_cleanup.volumes[0][0].title)

            service.activate_revision(document.id, original_revision.id)
            restored = service.get_directory(document.id)
            self.assertEqual("第七卷 雨夜", restored.volumes[0][0].title)
            self.assertEqual("雨夜", restored.volumes[0][1][0].title)
            self.assertNotEqual(original_revision.id, edit.revision.id)
            self.assertNotEqual(edit.revision.id, renamed_volume.revision.id)

    def test_merge_preserves_authoritative_volume_and_chapter_structure_and_exports(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            first_path = root / "first.txt"
            second_path = root / "second.txt"
            first_path.write_text(
                "第七卷 雨夜\n\n第787章 雨夜\n正文一。\n\n第788章 风声\n正文二。\n",
                encoding="utf-8",
            )
            second_path.write_text("第九章 归途\n正文三。\n", encoding="utf-8")
            service = DocumentLibraryService(initialized_database(root / "rusty.db"), root / "library")
            first = service.import_document(first_path).document
            second = service.import_document(second_path).document
            first_revision = service.list_revisions(first.id)[0]
            second_revision = service.list_revisions(second.id)[0]

            merged = service.merge_documents([first.id, second.id], "合并本")
            directory_result = service.get_directory(merged.id)
            merged_revision = service.list_revisions(merged.id)[0]
            txt_path = service.export_document(merged.id, "txt", root / "merged.txt")
            epub_path = service.export_document(merged.id, "epub", root / "merged.epub")

            self.assertEqual("merge", merged_revision.revision_type)
            self.assertEqual(["第七卷 雨夜"], [item[0].title for item in directory_result.volumes])
            self.assertEqual(
                ["雨夜", "风声"],
                [chapter.title for chapter in directory_result.volumes[0][1]],
            )
            self.assertEqual(["归途"], [chapter.title for chapter in directory_result.unassigned_chapters])
            self.assertEqual(first_revision.id, service.list_revisions(first.id)[0].id)
            self.assertEqual(second_revision.id, service.list_revisions(second.id)[0].id)

            exported_text = txt_path.read_text(encoding="utf-8")
            self.assertLess(exported_text.index("第七卷 雨夜"), exported_text.index("第一章 雨夜"))
            self.assertLess(exported_text.index("第二章 风声"), exported_text.index("第三章 归途"))
            self.assertEqual(1, exported_text.count("第七卷 雨夜"))
            self.assertEqual(1, exported_text.count("第一章 雨夜"))

            exported_book = epub.read_epub(str(epub_path))
            self.assertEqual(2, len(exported_book.toc))
            self.assertIsInstance(exported_book.toc[0], tuple)
            self.assertEqual(2, len(exported_book.toc[0][1]))


def _write_sample_epub(path: Path) -> None:
    book = epub.EpubBook()
    book.set_identifier("library-test")
    book.set_title("EPUB 示例")
    book.set_language("zh-CN")
    book.add_author("作者乙")
    chapter = epub.EpubHtml(title="开篇", file_name="chapter.xhtml", lang="zh-CN")
    chapter.content = "<h1>开篇</h1><p>第一段。</p>"
    book.add_item(chapter)
    book.toc = (chapter,)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)


if __name__ == "__main__":
    unittest.main()
