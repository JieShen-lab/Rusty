from __future__ import annotations

import sys
import sqlite3
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document

from rusty.importers.epub import parse_epub
from rusty.models import ExportPlanItem
from rusty.services import PipelineService, ProjectService


class ProjectServiceTests(unittest.TestCase):
    def test_delete_project_soft_deletes_from_project_list(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            txt_path = root / "delete-me.txt"
            database_path = root / "rusty.db"
            txt_path.write_text("1. Opening\nOriginal text.\n", encoding="utf-8")

            service = ProjectService(database_path)
            project_id = service.import_book(txt_path, root)
            chapters_before_delete = service.list_chapters(project_id)

            service.delete_project(project_id)
            projects_after_delete = service.list_projects()
            project_after_delete = service.get_project(project_id)
            chapters_after_delete = service.list_chapters(project_id)

        self.assertEqual(1, len(chapters_before_delete))
        self.assertEqual([], projects_after_delete)
        self.assertIsNone(project_after_delete)
        self.assertEqual(chapters_before_delete, chapters_after_delete)

    def test_save_chapter_rewrite_updates_export_text_and_can_clear(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            txt_path = root / "manual.txt"
            database_path = root / "rusty.db"
            export_path = root / "manual-export.txt"
            txt_path.write_text("1. Opening\nOriginal text.\n", encoding="utf-8")

            service = ProjectService(database_path)
            project_id = service.import_book(txt_path, root)
            chapter = service.list_chapters(project_id)[0]

            service.save_chapter_rewrite(chapter.id, "Manual rewrite.")
            updated = service.get_chapter(chapter.id)
            project_after_rewrite = service.get_project(project_id)
            unconfirmed_export = service.get_effective_export_chapters(project_id)[0]
            PipelineService(database_path).confirm_rewrite(chapter.id)
            service.export_txt(project_id, export_path)
            export_records = service.list_exports(project_id)
            exported_text = export_path.read_text(encoding="utf-8")
            connection = sqlite3.connect(database_path)
            try:
                rewrite_source, prompt_snapshot_json, anchor_snapshot_json = connection.execute(
                    """
                    SELECT rewrite_source, prompt_snapshot_json, anchor_snapshot_json
                    FROM chapter_rewrites
                    WHERE chapter_id = ?
                    """,
                    (chapter.id,),
                ).fetchone()
            finally:
                connection.close()

            service.save_chapter_rewrite(chapter.id, "")
            cleared = service.get_chapter(chapter.id)
            project_after_clear = service.get_project(project_id)

        self.assertIsNotNone(updated)
        self.assertEqual("Manual rewrite.", updated.rewritten_text)
        self.assertEqual("rewritten", updated.status)
        self.assertIsNone(unconfirmed_export.rewritten_text)
        self.assertEqual("original", unconfirmed_export.source_status)
        self.assertIsNotNone(project_after_rewrite)
        self.assertEqual(1, project_after_rewrite.completed_chapters)
        self.assertEqual(1, len(export_records))
        self.assertEqual("txt", export_records[0].export_format)
        self.assertEqual(str(export_path), export_records[0].output_path)
        self.assertEqual(14, export_records[0].word_count)
        self.assertEqual("manual", rewrite_source)
        self.assertIn("manual_edit", prompt_snapshot_json)
        self.assertEqual("{}", anchor_snapshot_json)
        self.assertIn("Manual rewrite.", exported_text)
        self.assertNotIn("Original text.", exported_text)
        self.assertIsNotNone(cleared)
        # Clearing is an append-only restore operation: the current projection
        # points at a new version containing the immutable original text.
        self.assertEqual("Original text.", cleared.rewritten_text)
        self.assertEqual("rewritten", cleared.status)
        self.assertIsNotNone(project_after_clear)
        self.assertEqual(1, project_after_clear.completed_chapters)

    def test_export_plan_reorders_renames_excludes_and_drives_exports(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            txt_path = root / "planned.txt"
            database_path = root / "rusty.db"
            txt_export_path = root / "planned-export.txt"
            epub_export_path = root / "planned-export.epub"
            txt_path.write_text(
                "1. One\nAlpha original.\n\n2. Two\nBeta original.\n\n3. Three\nGamma original.\n",
                encoding="utf-8",
            )

            service = ProjectService(database_path)
            project_id = service.import_book(txt_path, root)
            chapters = service.list_chapters(project_id)
            service.save_chapter_rewrite(chapters[1].id, "Manual beta rewrite.")
            PipelineService(database_path).confirm_rewrite(chapters[1].id)
            connection = sqlite3.connect(database_path)
            try:
                connection.execute(
                    "UPDATE chapters SET status = 'kept_original' WHERE id = ?",
                    (chapters[2].id,),
                )
                connection.commit()
            finally:
                connection.close()

            default_plan = service.list_export_plan(project_id)
            service.save_export_plan(
                project_id,
                [
                    ExportPlanItem(
                        chapter_id=chapters[1].id,
                        export_order=1,
                        export_title="Renamed Two",
                        include_in_export=True,
                    ),
                    ExportPlanItem(
                        chapter_id=chapters[0].id,
                        export_order=2,
                        export_title="Hidden One",
                        include_in_export=False,
                    ),
                    ExportPlanItem(
                        chapter_id=chapters[2].id,
                        export_order=3,
                        export_title="Renamed Three",
                        include_in_export=True,
                    ),
                ],
            )
            effective = service.get_effective_export_chapters(project_id)
            service.export_txt(project_id, txt_export_path)
            service.export_epub(project_id, epub_export_path)
            txt_export = txt_export_path.read_text(encoding="utf-8")
            parsed_epub = parse_epub(epub_export_path)
            export_records = service.list_exports(project_id)

        self.assertEqual(["1. One", "2. Two", "3. Three"], [item.export_title for item in default_plan])
        self.assertEqual(["Renamed Two", "Renamed Three"], [chapter.title for chapter in effective])
        self.assertEqual(["manual_rewrite", "kept_original"], [chapter.source_status for chapter in effective])
        self.assertIn("Renamed Two", txt_export)
        self.assertIn("Manual beta rewrite.", txt_export)
        self.assertIn("Renamed Three", txt_export)
        self.assertNotIn("Alpha original.", txt_export)
        self.assertEqual(["Renamed Two", "Renamed Three"], [chapter.title for chapter in parsed_epub.chapters])
        self.assertEqual(2, export_records[0].chapter_count)
        self.assertEqual(2, export_records[1].chapter_count)
        self.assertEqual(32, export_records[0].word_count)

    def test_save_export_plan_rejects_missing_or_foreign_chapters(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            first_path = root / "first.txt"
            second_path = root / "second.txt"
            database_path = root / "rusty.db"
            first_path.write_text("1. One\nAlpha.\n\n2. Two\nBeta.\n", encoding="utf-8")
            second_path.write_text("1. Other\nOther.\n", encoding="utf-8")

            service = ProjectService(database_path)
            first_project_id = service.import_book(first_path, root)
            second_project_id = service.import_book(second_path, root)
            first_chapters = service.list_chapters(first_project_id)
            foreign_chapter = service.list_chapters(second_project_id)[0]

            with self.assertRaises(ValueError):
                service.save_export_plan(
                    first_project_id,
                    [
                        ExportPlanItem(
                            chapter_id=first_chapters[0].id,
                            export_order=1,
                            export_title="Only one",
                            include_in_export=True,
                        )
                    ],
                )
            with self.assertRaises(ValueError):
                service.save_export_plan(
                    first_project_id,
                    [
                        ExportPlanItem(
                            chapter_id=first_chapters[0].id,
                            export_order=1,
                            export_title="One",
                            include_in_export=True,
                        ),
                        ExportPlanItem(
                            chapter_id=foreign_chapter.id,
                            export_order=2,
                            export_title="Foreign",
                            include_in_export=True,
                        ),
                    ],
                )

    def test_import_docx_persists_metadata_and_exports_epub(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            root = Path(directory)
            docx_path = root / "service.docx"
            database_path = root / "rusty.db"
            epub_path = root / "service.epub"

            document = Document()
            document.core_properties.title = "Service Book"
            document.core_properties.author = "Service Author"
            document.add_heading("Opening", level=1)
            document.add_paragraph("A saved chapter.")
            document.save(docx_path)

            service = ProjectService(database_path)
            project_id = service.import_book(docx_path, root)
            project = service.get_project(project_id)
            metadata = service.get_book_metadata(project_id)
            chapters = service.list_chapters(project_id)
            exported = service.export_epub(project_id, epub_path)
            parsed_export = parse_epub(exported)

        self.assertIsNotNone(project)
        self.assertEqual("Service Book", project.name)
        self.assertEqual("docx", project.source_format)
        self.assertEqual("Service Author", metadata["author"])
        self.assertEqual(1, len(chapters))
        self.assertEqual("Opening", chapters[0].title)
        self.assertEqual("Service Book", parsed_export.title)
        self.assertEqual("Service Author", parsed_export.author)


if __name__ == "__main__":
    unittest.main()
