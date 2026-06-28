from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document
from ebooklib import epub

from rusty.exporters.epub import export_epub
from rusty.importers.docx import parse_docx
from rusty.importers.epub import parse_epub
from rusty.models import ChapterRecord


class EpubDocxTests(unittest.TestCase):
    def test_parse_epub_reads_metadata_and_chapters(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            path = Path(directory) / "sample.epub"
            _write_sample_epub(path)

            parsed = parse_epub(path)

        self.assertEqual("Sample Book", parsed.title)
        self.assertEqual("Tester", parsed.author)
        self.assertEqual("en", parsed.language)
        self.assertEqual("epub", parsed.source_format)
        self.assertEqual(2, len(parsed.chapters))
        self.assertEqual("Chapter One", parsed.chapters[0].title)
        self.assertIn("First paragraph.", parsed.chapters[0].text)
        self.assertEqual("Chapter Two", parsed.chapters[1].title)

    def test_parse_docx_splits_heading_chapters(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            path = Path(directory) / "sample.docx"
            document = Document()
            document.core_properties.title = "Docx Book"
            document.core_properties.author = "Docx Author"
            document.add_heading("Part One", level=1)
            document.add_paragraph("Alpha line.")
            document.add_heading("Part Two", level=1)
            document.add_paragraph("Beta line.")
            document.save(path)

            parsed = parse_docx(path)

        self.assertEqual("Docx Book", parsed.title)
        self.assertEqual("Docx Author", parsed.author)
        self.assertEqual("docx", parsed.source_format)
        self.assertEqual(2, len(parsed.chapters))
        self.assertEqual("Part One", parsed.chapters[0].title)
        self.assertEqual("Alpha line.", parsed.chapters[0].text)
        self.assertEqual("Part Two", parsed.chapters[1].title)

    def test_export_epub_writes_readable_book(self) -> None:
        with tempfile.TemporaryDirectory(dir=Path.cwd()) as directory:
            path = Path(directory) / "export.epub"
            chapters = [
                ChapterRecord(
                    id=1,
                    project_id=1,
                    index=1,
                    title="One",
                    original_text="Alpha\nBeta",
                    rewritten_text=None,
                    word_count=9,
                    status="imported",
                    start_line=None,
                    end_line=None,
                )
            ]

            export_epub(chapters, path, title="Exported", author="Rusty", language="en")
            parsed = parse_epub(path)

        self.assertEqual("Exported", parsed.title)
        self.assertEqual("Rusty", parsed.author)
        self.assertEqual(1, len(parsed.chapters))
        self.assertEqual("One", parsed.chapters[0].title)
        self.assertIn("Alpha", parsed.chapters[0].text)


def _write_sample_epub(path: Path) -> None:
    book = epub.EpubBook()
    book.set_identifier("sample-id")
    book.set_title("Sample Book")
    book.set_language("en")
    book.add_author("Tester")

    chapter_one = epub.EpubHtml(title="Chapter One", file_name="chap_0001.xhtml", lang="en")
    chapter_one.content = "<h1>Chapter One</h1><p>First paragraph.</p>"
    chapter_two = epub.EpubHtml(title="Chapter Two", file_name="chap_0002.xhtml", lang="en")
    chapter_two.content = "<h1>Chapter Two</h1><p>Second paragraph.</p>"

    book.add_item(chapter_one)
    book.add_item(chapter_two)
    book.toc = (chapter_one, chapter_two)
    book.spine = ["nav", chapter_one, chapter_two]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)


if __name__ == "__main__":
    unittest.main()

